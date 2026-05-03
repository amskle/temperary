"""Tools: template parsing, filling, code analysis, report generation."""

import json
import logging
import os
import re
import subprocess
import tempfile
import time
from datetime import datetime
from typing import Any

from docx import Document

from config import config

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Template parsing — detects {{var}} style Jinja2 placeholders
# ---------------------------------------------------------------------------

_VARIABLE_PATTERN = re.compile(r"\{\{(\w+)\}\}")


def _extract_placeholders(text: str) -> list[dict]:
    """Extract placeholder info from text."""
    results: list[dict] = []
    for match in _VARIABLE_PATTERN.finditer(text):
        field_id = match.group(1)
        results.append({"field_id": field_id, "raw_placeholder": match.group(0)})
    return results


# Also support <var> style placeholders (angle brackets)
_VARIABLE_PATTERN_ANGLE = re.compile(r"<(\w+)>")


def _extract_placeholders_angle(text: str) -> list[dict]:
    """Extract angle-bracket placeholder info from text."""
    results: list[dict] = []
    for match in _VARIABLE_PATTERN_ANGLE.finditer(text):
        field_id = match.group(1)
        results.append({"field_id": field_id, "raw_placeholder": match.group(0)})
    return results


_ANALYZE_TEMPLATE_SYSTEM = (
    "You are an expert document analyst. Given the text content of a document "
    "template and user requirements, identify which paragraphs are section "
    "headers and which paragraphs need content to be filled in.\n\n"
    "Return ONLY a JSON object with this structure:\n"
    '{"sections": [\n'
    '  {"field_id": "purpose", "header_text": "一、实验目的", '
    '"content_index_after_header": 0, "is_complex": true},\n'
    '  {"field_id": "name", "header_text": "姓名", '
    '"content_index_after_header": 0, "is_complex": false}\n'
    "]}\n\n"
    "Rules:\n"
    "- field_id: a short English identifier for the section (snake_case)\n"
    "- header_text: the exact header text as it appears in the template. "
    "If a line starts with '[table X, row Y, col Z]', use ONLY the text "
    "AFTER that prefix as header_text (strip the prefix).\n"
    "- content_index_after_header: 0 means the paragraph immediately after "
    "the header, 1 means two paragraphs after, etc.\n"
    "- is_complex: true for academic content sections (purpose, principle, "
    "steps, results, analysis, conclusion), false for simple metadata "
    "(name, student ID, date, course)\n"
    "- Include ALL sections that need content, including metadata fields\n"
    "- If a paragraph already has placeholder text (like {{var}} or <var>), "
    "still include it — the placeholder indicates it needs filling\n"
    "- The header_text must be an EXACT substring match of the template text"
)


def analyze_template_with_llm(
    paragraphs: list[str],
    requirements: str,
) -> list[dict[str, Any]]:
    """Use LLM to analyze template structure and identify fillable sections.

    Returns a list of section dicts with keys:
        field_id, header_text, content_index_after_header, is_complex
    """
    # Deduplicate while preserving order (merged table cells repeat text)
    seen: set[str] = set()
    unique: list[str] = []
    for text in paragraphs:
        stripped = text.strip()
        if stripped and stripped not in seen:
            seen.add(stripped)
            unique.append(text)

    numbered = "\n".join(
        f"  [{i}] {text}" for i, text in enumerate(unique) if text.strip()
    )
    user_prompt = (
        f"--- Template paragraphs ---\n{numbered}\n\n"
        f"--- User requirements ---\n{requirements}\n\n"
        "Identify all sections that need content filled in. "
        "Return ONLY the JSON object."
    )

    raw = _call_llm(
        system_prompt=_ANALYZE_TEMPLATE_SYSTEM,
        user_prompt=user_prompt,
        temperature=0.1,
        max_tokens=8192,
        step_name="analyze_template",
    )

    parsed = _extract_json_from_llm_response(raw)
    sections = parsed.get("sections", [])

    # Validate structure
    valid = []
    for s in sections:
        if all(k in s for k in ("field_id", "header_text")):
            s.setdefault("content_index_after_header", 0)
            s.setdefault("is_complex", True)
            valid.append(s)

    logger.info(
        "Observation: LLM identified %d fillable section(s).",
        len(valid),
    )
    return valid


def parse_template(template_path: str, requirements: str = "") -> dict[str, Any]:
    """Parse a .docx template, returning immutable sections and variable fields.

    Returns
        {
          "immutable_sections": [{"type": str, "index": int, "text": str}, ...],
          "variable_fields": [{"field_id": str, "placeholder": str,
                               "context": str}, ...],
          "llm_sections": [{"field_id": str, "header_text": str, ...}] (optional)
        }

    Supports ``{{field}}`` and ``<field>`` placeholder styles.
    If no placeholders are found, falls back to LLM-based template analysis.
    """
    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    immutable_sections: list[dict[str, Any]] = []
    variable_fields: dict[str, dict[str, Any]] = {}
    paragraph_texts: list[str] = []

    # --- Paragraphs ---
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        paragraph_texts.append(para.text)  # keep original for LLM analysis
        if not text:
            continue

        # Try both {{var}} and <var> placeholder styles
        placeholders = _extract_placeholders(text) + _extract_placeholders_angle(text)
        if placeholders:
            for ph in placeholders:
                fid = ph["field_id"]
                if fid not in variable_fields:
                    variable_fields[fid] = {
                        "field_id": fid,
                        "placeholder": ph["raw_placeholder"],
                        "context": text,
                    }

        immutable_sections.append({"type": "paragraph", "index": i, "text": text})

    # --- Tables ---
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue

                placeholders = (
                    _extract_placeholders(text) + _extract_placeholders_angle(text)
                )
                if placeholders:
                    for ph in placeholders:
                        fid = ph["field_id"]
                        if fid not in variable_fields:
                            variable_fields[fid] = {
                                "field_id": fid,
                                "placeholder": ph["raw_placeholder"],
                                "context": (
                                    f"[table {ti}, row {ri}, col {ci}] {text}"
                                ),
                            }

                # Collect table cell text for LLM analysis (with location tag)
                for para in cell.paragraphs:
                    if para.text.strip():
                        paragraph_texts.append(
                            f"[table {ti}, row {ri}, col {ci}] {para.text}"
                        )

                immutable_sections.append(
                    {
                        "type": "table_cell",
                        "table_index": ti,
                        "row": ri,
                        "col": ci,
                        "text": text,
                    }
                )

    result: dict[str, Any] = {
        "immutable_sections": immutable_sections,
        "variable_fields": list(variable_fields.values()),
    }

    # Fallback: if no placeholders found, use LLM to analyze template structure
    if not variable_fields and requirements:
        logger.info(
            "  Observation: No placeholders found. "
            "Falling back to LLM template analysis."
        )
        llm_sections = analyze_template_with_llm(paragraph_texts, requirements)
        if llm_sections:
            result["llm_sections"] = llm_sections
            # Also populate variable_fields for compatibility with downstream code
            for s in llm_sections:
                result["variable_fields"].append({
                    "field_id": s["field_id"],
                    "placeholder": "",
                    "context": s.get("header_text", ""),
                })

    return result



# ---------------------------------------------------------------------------
# Template filling — DocxTemplate renders Jinja2 {{field}} placeholders at XML level
# ---------------------------------------------------------------------------


def fill_template(
    template_path: str,
    content_mapping: dict[str, str],
    output_path: str | None = None,
) -> str:
    """Render ``{{field}}`` placeholders with DocxTemplate and save.

    DocxTemplate works at the XML level, so it naturally handles
    placeholders split across multiple runs by Word.
    """
    from docxtpl import DocxTemplate

    doc = DocxTemplate(template_path)
    doc.render(content_mapping)

    if output_path is None:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(config.output_dir, f"output_{ts}.docx")

    doc.save(output_path)
    return os.path.abspath(output_path)


def _find_header_in_paragraphs(
    paragraphs, header_text: str
) -> int | None:
    """Find header_text in doc.paragraphs by substring match. Returns index or None."""
    for i, para in enumerate(paragraphs):
        if header_text in para.text:
            return i
    return None


def _write_to_paragraph(paragraphs, content_idx: int, content: str) -> bool:
    """Write content into paragraphs[content_idx], preserving style. Returns success."""
    if content_idx >= len(paragraphs):
        return False
    target_para = paragraphs[content_idx]
    style = target_para.style
    for run in target_para.runs:
        run.text = ""
    if target_para.runs:
        target_para.runs[0].text = content
    else:
        target_para.add_run(content)
    target_para.style = style
    return True


def _find_header_in_tables(
    tables, header_text: str
) -> tuple | None:
    """Find header_text in table cells by substring match.

    Returns (table, row_idx, content_cell, content_para_idx) or None.

    For merged cells (section headers spanning the full row), content goes
    into the same cell at paragraph index 1.
    For non-merged cells with an empty neighbour (metadata fields like
    姓名/班级/学号), content goes into that neighbour cell at paragraph 0.
    """
    for table in tables:
        for ri, row in enumerate(table.rows):
            cells = list(row.cells)
            for ci, cell in enumerate(cells):
                if header_text not in cell.text:
                    continue

                # Check if this cell is horizontally merged with the next one
                if ci + 1 < len(cells) and cells[ci + 1]._tc is cell._tc:
                    # Merged cell — content goes at paragraph[1] in the same cell
                    return (table, ri, cell, 1)

                # Non-merged: look for an adjacent empty cell for content
                if ci + 1 < len(cells) and not cells[ci + 1].text.strip():
                    return (table, ri, cells[ci + 1], 0)

                # Fallback: write below the header in the same cell
                return (table, ri, cell, 1)
    return None


def _write_to_table_cell(cell, content_para_idx: int, content: str) -> bool:
    """Write content into a table cell at the given paragraph index.

    The header text is typically at p[0]; content goes at p[content_para_idx].
    If the paragraph doesn't exist, adds new paragraphs as needed.
    Handles merged cells where ``add_paragraph`` may not persist by falling
    back to direct XML manipulation.
    """
    # Ensure enough paragraphs exist in the cell
    while len(cell.paragraphs) <= content_para_idx:
        added = cell.add_paragraph("")
        # If add_paragraph didn't actually increase the count (merged-cell
        # edge case in some python-docx versions), insert via raw XML.
        if len(cell.paragraphs) <= content_para_idx:
            from docx.oxml import OxmlElement
            new_p = OxmlElement("w:p")
            cell._tc.append(new_p)

    target_para = cell.paragraphs[content_para_idx]
    style = target_para.style

    # Clear existing runs
    for run in target_para.runs:
        run.text = ""
    if target_para.runs:
        target_para.runs[0].text = content
    else:
        target_para.add_run(content)
    target_para.style = style
    return True


def fill_template_intelligent(
    template_path: str,
    llm_sections: list[dict[str, Any]],
    content_mapping: dict[str, str],
    output_path: str | None = None,
) -> str:
    """Fill template by locating content via header text matching.

    Supports two template layouts:
    1. Paragraph-based: headers and content in doc.paragraphs (original behavior)
    2. Table-based: headers and content inside table cells (new support)

    For each section, tries paragraph search first, then falls back to table search.
    """
    doc = Document(template_path)
    paragraphs = doc.paragraphs

    filled_count = 0
    for section in llm_sections:
        header_text = section.get("header_text", "")
        field_id = section.get("field_id", "")
        offset = section.get("content_index_after_header", 0)

        if field_id not in content_mapping or not content_mapping[field_id].strip():
            logger.warning(
                "  Skipping section '%s' — no content generated.", field_id,
            )
            continue

        content = content_mapping[field_id]

        # Strategy 1: search in doc.paragraphs (original logic)
        header_idx = _find_header_in_paragraphs(paragraphs, header_text)
        if header_idx is not None:
            content_idx = header_idx + 1 + offset
            if _write_to_paragraph(paragraphs, content_idx, content):
                filled_count += 1
                logger.info(
                    "  Filled section '%s' at paragraph %d (%d chars).",
                    field_id, content_idx, len(content),
                )
            else:
                logger.warning(
                    "  Content paragraph index %d out of range for section '%s'.",
                    content_idx, field_id,
                )
            continue

        # Strategy 2: search in table cells (new support for table-based templates)
        found = _find_header_in_tables(doc.tables, header_text)
        if found:
            _table, _row_idx, cell, base_para_idx = found
            content_para_idx = base_para_idx + offset
            if _write_to_table_cell(cell, content_para_idx, content):
                filled_count += 1
                logger.info(
                    "  Filled section '%s' in table cell (%d chars).",
                    field_id, len(content),
                )
            else:
                logger.warning(
                    "  Failed to write to table cell for section '%s'.",
                    field_id,
                )
            continue

        logger.warning(
            "  Could not find header '%s' in template. Skipping section '%s'.",
            header_text, field_id,
        )

    logger.info("  Observation: Filled %d section(s) via intelligent filling.", filled_count)

    if output_path is None:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(config.output_dir, f"output_{ts}.docx")

    doc.save(output_path)
    return os.path.abspath(output_path)


# ---------------------------------------------------------------------------
# Xiaomi MiMo LLM call (with retry & fallback) — ReAct-style observability
# ---------------------------------------------------------------------------


def _call_llm(
    system_prompt: str,
    user_prompt: str,
    temperature: float = 0.3,
    max_tokens: int = 8192,
    step_name: str = "llm_call",
) -> str:
    """Call Xiaomi MiMo (Anthropic-compatible) with retry + fallback + ReAct logging."""
    import anthropic

    models_to_try = [config.xiaomi_model, config.xiaomi_fallback_model]
    # deduplicate while preserving order
    models_to_try = list(dict.fromkeys(models_to_try))

    last_exc: Exception | None = None

    logger.info("Thought: Need to call LLM (model=%s)", models_to_try[0])
    logger.info("Action: _call_llm[%s]", step_name)

    if not config.xiaomi_api_key:
        raise RuntimeError(
            "XIAOMI_API_KEY is not set. "
            "Export it: export XIAOMI_API_KEY=your-key"
        )

    for attempt in range(config.max_retries):
        model = models_to_try[0] if attempt == 0 else models_to_try[-1]
        try:
            client = anthropic.Anthropic(
                api_key=config.xiaomi_api_key,
                base_url=config.xiaomi_base_url,
                timeout=config.request_timeout,
            )
            resp = client.messages.create(
                model=model,
                max_tokens=max_tokens,
                temperature=temperature,
                system=system_prompt,
                messages=[
                    {"role": "user", "content": user_prompt},
                ],
            )
            # MiMo (reasoning model) may return ThinkingBlock + TextBlock.
            # Prefer TextBlock.text, fall back to ThinkingBlock.thinking.
            content = ""
            for block in resp.content:
                block_type = type(block).__name__
                has_text = hasattr(block, "text") and bool(block.text)
                has_thinking = hasattr(block, "thinking") and bool(block.thinking)
                logger.info(
                    "  [LLM block] type=%s has_text=%s has_thinking=%s",
                    block_type, has_text, has_thinking,
                )
                if has_text:
                    content = block.text
                    break
            if not content:
                for block in resp.content:
                    if hasattr(block, "thinking") and block.thinking:
                        content = block.thinking
                        break
            logger.info(
                "Observation: LLM responded successfully (%d chars, model=%s)",
                len(content), model,
            )
            if content:
                logger.info("  LLM content preview: %s…", content[:300])
            return content

        except Exception as exc:
            last_exc = exc
            logger.warning(
                "LLM call failed (attempt %d, model=%s): %s",
                attempt + 1, model, exc,
            )
            if attempt < config.max_retries - 1:
                delay = config.retry_base_delay * (2**attempt)
                logger.info(
                    "Observation: will retry in %.0fs (exponential backoff) …",
                    delay,
                )
                time.sleep(delay)

    raise RuntimeError(
        f"LLM call failed after {config.max_retries} retries. "
        f"Last error: {last_exc}"
    )


# ---------------------------------------------------------------------------
# Code analysis
# ---------------------------------------------------------------------------

_ANALYZE_CODE_SYSTEM = (
    "You are an expert programming tutor. Given source code, "
    "explain: 1) what the code does, "
    "2) its key algorithms / logic, "
    "3) what output or result it produces. "
    "Be concise but thorough."
)


def analyze_code(code: str) -> str:
    """Analyze source code and return a structured explanation."""
    prompt = (
        "Analyze the following code and explain:\n"
        "- Purpose and functionality\n"
        "- Key algorithms and data structures\n"
        "- Expected output / results\n\n"
        f"```\n{code}\n```"
    )
    return _call_llm(
        system_prompt=_ANALYZE_CODE_SYSTEM,
        user_prompt=prompt,
        temperature=0.2,
        step_name="analyze_code",
    )


# ---------------------------------------------------------------------------
# Dynamic code execution sandbox
# ---------------------------------------------------------------------------


def execute_code_sandbox(code: str, timeout: int | None = None) -> dict[str, str]:
    """Execute Python code in a sandboxed subprocess and capture actual output.

    Returns dict with keys: stdout, stderr, exit_code, error.
    The captured output is used to ground the LLM's result/analysis in real data.
    """
    if timeout is None:
        timeout = config.code_execution_timeout

    result: dict[str, str] = {
        "stdout": "",
        "stderr": "",
        "exit_code": "-1",
        "error": "",
    }

    tmp_path = ""
    try:
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".py", delete=False, encoding="utf-8"
        ) as f:
            f.write(code)
            tmp_path = f.name

        proc = subprocess.run(
            ["python3", tmp_path],
            capture_output=True,
            text=True,
            timeout=timeout,
            cwd=tempfile.gettempdir(),
        )
        max_chars = config.max_code_output_chars
        result["stdout"] = proc.stdout[:max_chars]
        result["stderr"] = proc.stderr[:max_chars]
        result["exit_code"] = str(proc.returncode)

    except subprocess.TimeoutExpired:
        result["error"] = f"Code execution timed out after {timeout}s"
    except Exception as exc:
        result["error"] = str(exc)
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    return result


# ---------------------------------------------------------------------------
# Report content generation
# ---------------------------------------------------------------------------

_GENERATE_SINGLE_FIELD_SYSTEM = (
    "You are an AI assistant helping to write ONE specific section of a college "
    "lab report. Generate detailed, academically rigorous content for this "
    "section ONLY. The content must be substantive — include theoretical "
    "explanations, quantitative analysis, and domain-specific terminology "
    "appropriate for college-level coursework.\n\n"
    "CRITICAL RULES:\n"
    "- Write AT LEAST {min_chars} characters of content (Chinese characters count individually).\n"
    "- Write 3-8 substantial paragraphs with complete sentences.\n"
    "- Include specific details, equations, or data references where relevant.\n"
    "- Reference code execution output or experimental data when provided.\n"
    "- NEVER output placeholders, brackets, or generic fallback text.\n"
    "- NEVER write '[Content for...]' or any similar placeholder.\n"
    "- Return ONLY the raw content text — no JSON wrapper, no field ID label."
)

_GENERATE_SIMPLE_FIELDS_SYSTEM = (
    "You are an AI assistant filling in simple metadata fields for a lab report. "
    "Generate appropriate values for fields like student name, ID, date, etc.\n"
    "Return ONLY a JSON object with field_id as key and value as content."
)


def _extract_total_char_target(requirements: str, num_complex_fields: int) -> int:
    """Extract total character count target from user requirements.

    Looks for patterns like '3000字', '3000字以上', '不少于3000字', '3000+字',
    '至少3000字符' etc. Returns a per-field target (total / num_fields),
    or 0 if no count requirement found.
    """
    if num_complex_fields <= 0:
        return 0

    # Match patterns like: 3000字, 3000字以上, 不少于3000字, 至少3000字, 3000+字
    patterns = [
        r"(?:不少于|至少|最少|不低于|>\s*)(\d+)\s*(?:字|字符|words?|chars?)",
        r"(\d{3,})\s*(?:字|字符)\s*(?:以上|及以上|以上的内容|以上的内容)",
        r"(\d{3,})\s*\+?\s*(?:字|字符|words?|chars?)",
        r"(?:字数|字符数|篇幅)\s*(?:要求|限制|：|:)?\s*(?:不少于|至少|>=?\s*)?(\d{3,})",
    ]
    for pattern in patterns:
        match = re.search(pattern, requirements, re.IGNORECASE)
        if match:
            total = int(match.group(1))
            per_field = total // num_complex_fields
            logger.info(
                "  Detected char target: %d total / %d complex fields = %d per field",
                total, num_complex_fields, per_field,
            )
            return per_field

    return 0

def _build_few_shot_text(few_shot_examples: list[dict] | None) -> str:
    """Format few-shot examples into a prompt-ready text block."""
    if not few_shot_examples:
        return ""
    parts = []
    for i, ex in enumerate(few_shot_examples, 1):
        parts.append(
            f"Example {i}:\n  Requirement: {ex['requirement']}\n"
            f"  Content: {ex['content']}"
        )
    return "--- Similar past reports ---\n" + "\n\n".join(parts)


def _build_immutable_context(template_info: dict[str, Any]) -> str:
    """Build the immutable sections context string from template info."""
    return "\n".join(
        f"  [{s['type']}] {s['text'][:3000]}"
        for s in template_info.get("immutable_sections", [])
    )


def _extract_json_from_llm_response(raw: str) -> dict:
    """Extract and parse a JSON object from LLM output.

    Handles markdown code fences (`` ```json ... ``` ``) and stray text
    surrounding the JSON object. Returns an empty dict on failure.
    """
    # Strip markdown code fences — use greedy match to capture full nested JSON
    fence_match = re.search(r"```(?:json)?\s*(\{.*\})\s*```", raw, re.DOTALL)
    if fence_match:
        raw = fence_match.group(1)

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass

    # Fallback: find the outermost { ... } pair (handles nested JSON correctly)
    brace_start = raw.find("{")
    brace_end = raw.rfind("}")
    if brace_start != -1 and brace_end > brace_start:
        try:
            return json.loads(raw[brace_start : brace_end + 1])
        except json.JSONDecodeError:
            pass

    logger.warning("  Failed to extract JSON from LLM response (%d chars). Preview: %s…",
                   len(raw), raw[:200])
    return {}


def _generate_simple_fields(
    simple_fields: list[dict],
    immutable_context: str,
    report_requirements: str,
    few_shot_text: str,
) -> dict[str, str]:
    """Generate simple metadata fields (name, student_id, date) in a single batch."""
    if not simple_fields:
        return {}

    fields_desc = "\n".join(
        f"  - {f['field_id']}  (context: {f['context'][:3000]})"
        for f in simple_fields
    )
    user_prompt = (
        f"--- Template structure ---\n{immutable_context}\n\n"
        f"--- Simple fields to fill ---\n{fields_desc}\n\n"
        f"--- User requirements ---\n{report_requirements}\n\n"
        f"{few_shot_text}\n\n"
        "Return a JSON object with field_id as key and content as value."
    )

    raw = _call_llm(
        system_prompt=_GENERATE_SIMPLE_FIELDS_SYSTEM,
        user_prompt=user_prompt,
        temperature=0.3,
        max_tokens=8192,
        step_name="generate_simple_fields",
    )

    result = _extract_json_from_llm_response(raw)

    # Ensure all requested fields are present
    out: dict[str, str] = {}
    for f in simple_fields:
        fid = f["field_id"]
        out[fid] = result.get(fid, "").strip() if result.get(fid) else ""

    return out


_SUMMARIZE_PRIOR_SYSTEM = (
    "You are a precise academic summarizer. Condense the following generated "
    "lab report sections into a concise, structured summary. Preserve key "
    "details: experimental values, equations, named methods, conclusions, "
    "and cross-references. The summary will be used as context for generating "
    "the NEXT section, so include anything needed for logical continuity."
)


def _summarize_prior_content(prior_text: str) -> str:
    """Compress long prior-content text into a summary via an LLM call."""
    raw = _call_llm(
        system_prompt=_SUMMARIZE_PRIOR_SYSTEM,
        user_prompt=(
            "Summarize the following lab report sections. "
            "Keep all quantitative data, named methods, equations, and key "
            "conclusions. Write a single flowing paragraph (≤300 words).\n\n"
            f"{prior_text}"
        ),
        temperature=0.2,
        max_tokens=8192,
        step_name="summarize_prior",
    )
    return raw.strip()


def _build_prior_context(prior_content: dict[str, str]) -> str:
    """Build prior-section context, summarising if the raw text is too long."""
    if not prior_content:
        return ""

    prior_lines = []
    for pk, pv in prior_content.items():
        preview = pv[:500].replace("\n", " ")
        prior_lines.append(f"  [{pk}]: {preview}")

    raw_text = (
        "--- Previously generated sections (read these to maintain "
        "consistency and avoid contradiction) ---\n"
        + "\n".join(prior_lines)
    )

    if len(raw_text) > config.prior_content_summary_threshold:
        logger.info(
            "  Prior content length %d exceeds threshold %d — summarising.",
            len(raw_text), config.prior_content_summary_threshold,
        )
        summary = _summarize_prior_content(raw_text)
        return (
            "--- Summary of previously generated sections (read for "
            "consistency and logical continuity) ---\n"
            f"{summary}\n\n"
        )

    return raw_text + "\n\n"


def _generate_single_field(
    field: dict,
    immutable_context: str,
    report_requirements: str,
    code_analysis: str,
    code_execution_output: str,
    few_shot_text: str,
    prior_content: dict[str, str],
    retry_feedback: str = "",
    min_chars: int = 0,
) -> str:
    """Generate one complex field with full context including prior sections."""
    fid = field["field_id"]
    context = field.get("context", "")[:3000]

    prior_text = _build_prior_context(prior_content)

    code_section = ""
    if code_analysis:
        code_section += f"--- Code analysis ---\n{code_analysis}\n\n"
    if code_execution_output and code_execution_output.strip():
        code_section += (
            f"--- Actual code execution output (ground truth) ---\n"
            f"{code_execution_output}\n\n"
        )

    retry_section = ""
    if retry_feedback:
        retry_section = (
            "--- IMPORTANT: Previous generation attempt failed ---\n"
            f"Issues to fix: {retry_feedback}\n"
            "Please ensure the content above is substantive, complete, and "
            "free of placeholder or generic fallback text.\n\n"
        )

    # Use per-field char target if available, otherwise default
    effective_min = max(min_chars, 500)
    char_instruction = (
        f"IMPORTANT: This section MUST be at least {effective_min} characters "
        f"(Chinese characters count individually). Write enough detail to meet "
        f"this requirement. "
    )

    user_prompt = (
        f"--- Template structure (for context) ---\n{immutable_context}\n\n"
        f"--- Field to generate ---\n"
        f"Field ID: {fid}\n"
        f"Template context: {context}\n\n"
        f"{prior_text}"
        f"--- User requirements ---\n{report_requirements}\n\n"
        f"{code_section}"
        f"{retry_section}"
        f"{few_shot_text}\n\n"
        f"{char_instruction}"
        f"Generate the complete content for the \"{fid}\" section. "
        f"Write detailed paragraphs with substantive academic content. "
        f"Return ONLY the raw content — no labels, no JSON."
    )

    raw = _call_llm(
        system_prompt=_GENERATE_SINGLE_FIELD_SYSTEM.format(min_chars=effective_min),
        user_prompt=user_prompt,
        temperature=0.4,
        max_tokens=8192,
        step_name=f"generate_field_{fid}",
    )

    return raw.strip()


def generate_report_iterative(
    report_requirements: str,
    code_analysis: str,
    code_execution_output: str,
    template_info: dict[str, Any],
    few_shot_examples: list[dict] | None = None,
    retry_feedback: str = "",
) -> dict[str, str]:
    """Generate report content iteratively — one complex field at a time.

    Simple fields (name, student_id, date) are batched. Complex fields
    (purpose, principle, steps, result, analysis, conclusion) are generated
    individually with all prior content passed as context. This forces the
    LLM to focus its full attention and output tokens on one section at a time.

    If *retry_feedback* is non-empty, it is injected into each complex-field
    prompt so the LLM knows what failed the previous validation pass.
    """
    variable_fields = template_info.get("variable_fields", [])

    # Separate simple vs complex, preserving template order
    simple_fields = [
        f for f in variable_fields if f["field_id"] not in config.complex_field_ids
    ]
    complex_fields = [
        f for f in variable_fields if f["field_id"] in config.complex_field_ids
    ]

    immutable_context = _build_immutable_context(template_info)
    few_shot_text = _build_few_shot_text(few_shot_examples)

    # Calculate per-field character target from user requirements
    min_chars_per_field = _extract_total_char_target(
        report_requirements, len(complex_fields)
    )

    result: dict[str, str] = {}

    # Phase 1: batch-generate simple metadata fields
    if simple_fields:
        simple_result = _generate_simple_fields(
            simple_fields, immutable_context, report_requirements, few_shot_text
        )
        result.update(simple_result)

    # Phase 2: generate each complex field individually with chained context
    for field in complex_fields:
        fid = field["field_id"]
        logger.info("Action: Generating field '%s' (iterative) …", fid)
        generated = _generate_single_field(
            field=field,
            immutable_context=immutable_context,
            report_requirements=report_requirements,
            code_analysis=code_analysis,
            code_execution_output=code_execution_output,
            few_shot_text=few_shot_text,
            prior_content=result,
            retry_feedback=retry_feedback,
            min_chars=min_chars_per_field,
        )
        result[fid] = generated
        logger.info("Observation: '%s' generated (%d chars)", fid, len(generated))

    # Fill any missing fields
    for field in variable_fields:
        fid = field["field_id"]
        if fid not in result or not result[fid].strip():
            result[fid] = f"[Content for {fid} — see generated output below]"

    return result
