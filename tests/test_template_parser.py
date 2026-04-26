"""Unit tests for template parsing."""

import os
import tempfile

from docx import Document

import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from tools import parse_template


def _create_docx_with_fields(path: str) -> None:
    """Helper: create a minimal .docx with ``{{var}}`` placeholders."""
    doc = Document()
    doc.add_paragraph("实验报告")
    doc.add_paragraph("姓名：{{name}}    学号：{{student_id}}")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("")
    doc.add_paragraph("二、实验原理")
    doc.add_paragraph("{{principle}}")
    doc.add_paragraph("三、实验结果")
    doc.add_paragraph("{{result}}")

    # Table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "测量项目"
    table.cell(0, 1).text = "{{value}}"
    table.cell(1, 0).text = "理论值"
    table.cell(1, 1).text = "{{expected}}"
    table.cell(2, 0).text = "结论"
    table.cell(2, 1).text = "{{conclusion}}"

    doc.save(path)


def _create_docx_no_fields(path: str) -> None:
    """Helper: create a .docx with NO variable fields."""
    doc = Document()
    doc.add_paragraph("固定标题")
    doc.add_paragraph("这是一段不可变的内容。")
    doc.add_paragraph("姓名：张三")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    doc.save(path)


class TestParseTemplate:
    def test_identifies_variable_fields(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            _create_docx_with_fields(path)
            result = parse_template(path)

            field_ids = {f["field_id"] for f in result["variable_fields"]}
            expected = {
                "name", "student_id", "principle", "result",
                "value", "expected", "conclusion",
            }
            assert field_ids == expected, f"Missing fields: {expected - field_ids}"

            # Verify placeholder format is {{field_id}}
            for f in result["variable_fields"]:
                assert f["placeholder"] == "{{" + f["field_id"] + "}}"
        finally:
            os.unlink(path)

    def test_immutable_sections_contain_paragraphs(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            _create_docx_with_fields(path)
            result = parse_template(path)

            texts = [s["text"] for s in result["immutable_sections"]]
            assert any("实验报告" in t for t in texts)
            assert any("实验目的" in t for t in texts)
            assert any("实验原理" in t for t in texts)
            assert any("测量项目" in t for t in texts)
        finally:
            os.unlink(path)

    def test_no_variable_fields(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            _create_docx_no_fields(path)
            result = parse_template(path)
            assert len(result["variable_fields"]) == 0
            assert len(result["immutable_sections"]) > 0
        finally:
            os.unlink(path)

    def test_raises_on_missing_file(self):
        import pytest

        with pytest.raises(FileNotFoundError):
            parse_template("/nonexistent/path.docx")
